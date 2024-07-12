import pandas as pd
from docx import Document
from docx.oxml.ns import qn
import numpy

# Load the Excel file
excel_path = './data/工作簿1.xlsx'
df = pd.read_excel(excel_path)

# Create a new Word document
doc = Document()

# Define a function to add sections to the Word document
def add_section(doc, section_number, section_title, content):
    doc.add_heading(f"{section_number} {section_title}", level=1)
    sub_section_number = 1
    sub_sub_section_number = 1
    sub_sub_sub_section_number = 1
    sub_sub_sub_sub_section_number = 1
    current_subtitle1 = None
    current_subtitle2 = None
    current_subtitle3 = None
    current_subtitle4 = None
    descriptions = []

    for idx, row in content.iterrows():
        subtitle1 = row['客户需求']
        subtitle2 = row.get('二级模块')
        subtitle3 = row.get('功能过程')
        description = row['子过程描述']

        # If subtitle1 is not NaN, update the current subtitle1 and reset subtitle2
        if pd.notna(subtitle1):
            current_subtitle1 = subtitle1
            if descriptions:
                combined_description = '，'.join(descriptions) + '。'
                doc.add_paragraph(combined_description)
            doc.add_heading(f"{current_subtitle1}", level=2)
            sub_section_number += 1
            sub_sub_section_number = 1  # Reset sub_sub_section_number when subtitle1 changes
            sub_sub_sub_section_number = 1  # Reset sub_sub_sub_section_number when subtitle1 changes
            current_subtitle2 = None  # Reset subtitle2 when subtitle1 changes
            current_subtitle3 = None  # Reset subtitle3 when subtitle1 changes
            current_subtitle4 = None  # Reset subtitle4 when subtitle1 changes
            descriptions = []

        # If subtitle2 is not NaN, update the current subtitle2 and add a heading
        if pd.notna(subtitle2):
            current_subtitle2 = subtitle2
            if descriptions:
                combined_description = '，'.join(descriptions) + '。'
                doc.add_paragraph(combined_description)
            doc.add_heading(f"{current_subtitle2}", level=3)
            sub_sub_section_number += 1
            sub_sub_sub_section_number = 1
            descriptions = []

        # If subtitle3 is not NaN, update the current subtitle3 and add a heading
        if pd.notna(subtitle3):
            current_subtitle3 = subtitle3
            if descriptions:
                combined_description = '，'.join(descriptions) + '。'
                doc.add_paragraph(combined_description)
            # doc.add_heading(f"{section_number}.{sub_section_number-1}.{sub_sub_section_number-1}. {current_subtitle3}", level=4)
            p = doc.add_paragraph()
            run = p.add_run(subtitle3)
            run.bold = True
            descriptions = []

         # Add the description under the current subtitle
        if pd.notna(description):
            descriptions.append(description)

    # Add remaining descriptions at the end of the section
    if descriptions:
        combined_description = '，'.join(descriptions) + '。'
        doc.add_paragraph(combined_description)

# Iterate through the Excel data and populate the Word document
section_title = "业务功能需求"
section_number = ""  # Changeable

add_section(doc, section_number, section_title, df)

doc.styles['Normal'].font.name = 'Times New Roman'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

# Save the Word document
word_path = './data/Generated_Document.docx'
doc.save(word_path)
print(f"Word document saved at {word_path}")
